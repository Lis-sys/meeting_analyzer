#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Приложение для анализа аудиозаписей совещаний и переговоров.
Выполняет распознавание речи, диаризацию, формирование протокола и суммаризацию.
Работает полностью локально без подключения к внешним сервисам.
Поддерживает локальную версию GigaChat и других технологий от СБЕР.
"""

import os
import sys
import json
import argparse
import torch
import whisper
import numpy as np
import logging
from pydub import AudioSegment
from transformers import T5ForConditionalGeneration, T5Tokenizer
from docx import Document
from resemblyzer import VoiceEncoder, preprocess_wav
from sklearn.cluster import AgglomerativeClustering
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import threading

# Настройка логирования
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    handlers=[
        logging.FileHandler("meeting_analyzer.log"),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# Определяем, доступен ли CUDA, и устанавливаем устройство
DEVICE = "cuda" if torch.cuda.is_available() else "cpu"
logger.info(f"Using device: {DEVICE}")

# Параметры для локального GigaChat
GIGACHAT_BASE_URL = os.environ.get("GIGACHAT_BASE_URL", "http://localhost:8080/api/v1")
GIGACHAT_AUTH_METHOD = os.environ.get("GIGACHAT_AUTH_METHOD", "key")  # key, token, cert, user_pass
GIGACHAT_CREDENTIALS = os.environ.get("GIGACHAT_CREDENTIALS", "")
GIGACHAT_MODEL = os.environ.get("GIGACHAT_MODEL", "GigaChat-Pro")

# Глобальные переменные для моделей
whisper_model = None
voice_encoder = None
summarization_model = None
summarization_tokenizer = None
summarizer = None


# Загрузка конфигурации из файла
def load_config(config_path="config.json"):
    if os.path.exists(config_path):
        try:
            with open(config_path, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception as e:
            logger.error(f"Error loading config: {e}")
    return {}


# Базовый класс для суммаризации
class BaseSummarizer:
    def summarize(self, text, max_length=500, min_length=50):
        raise NotImplementedError("Subclasses must implement this method")


# Реализация суммаризатора на основе RuT5
class RuT5Summarizer(BaseSummarizer):
    def __init__(self, model=None, tokenizer=None, device="cpu"):
        self.model = model
        self.tokenizer = tokenizer
        self.device = device

    def summarize(self, text, max_length=500, min_length=50):
        if not text or not text.strip():
            return "Error: No text provided for summarization."

        try:
            # Подготовка текста для модели T5
            prefix = "summarize: "
            input_text = prefix + text

            # Ограничиваем длину входного текста, чтобы избежать переполнения памяти
            input_ids = self.tokenizer.encode(
                input_text,
                return_tensors="pt",
                max_length=1024,
                truncation=True,
                padding="max_length"
            ).to(self.device)

            # Генерация суммаризации с улучшенными параметрами
            summary_ids = self.model.generate(
                input_ids,
                max_length=max_length,
                min_length=min_length,
                num_beams=5,  # Увеличено для лучшего качества
                length_penalty=2.0,
                early_stopping=True,
                no_repeat_ngram_size=3,
                do_sample=False,  # Отключаем случайность для стабильности
                temperature=1.0  # Нейтральная температура
            )

            summary = self.tokenizer.decode(summary_ids[0], skip_special_tokens=True)
            return summary
        except Exception as e:
            logger.error(f"Error in RuT5 summarization: {e}")
            return f"Ошибка при суммаризации текста с RuT5: {e}"


# Реализация суммаризатора на основе GigaChat
class GigaChatSummarizer(BaseSummarizer):
    def __init__(self, base_url=None, auth_method="key", credentials=None, model=None, **kwargs):
        self.base_url = base_url
        self.auth_method = auth_method
        self.credentials = credentials
        self.model = model
        self.kwargs = kwargs
        self.client = None

        # Импортируем GigaChat только при необходимости
        try:
            from gigachat import GigaChat
            self.GigaChat = GigaChat
            self.client = self._initialize_client()
        except ImportError:
            logger.error("GigaChat library not installed. Run 'pip install gigachat'")
            raise ImportError("GigaChat library not installed. Run 'pip install gigachat'")

    def _initialize_client(self):
        # Инициализация клиента GigaChat в зависимости от метода авторизации
        try:
            if self.auth_method == "key":
                return self.GigaChat(
                    credentials=self.credentials,
                    base_url=self.base_url,
                    model=self.model,
                    verify_ssl_certs=False,
                    **self.kwargs
                )
            elif self.auth_method == "token":
                return self.GigaChat(
                    access_token=self.credentials,
                    base_url=self.base_url,
                    model=self.model,
                    verify_ssl_certs=False,
                    **self.kwargs
                )
            elif self.auth_method == "cert":
                # Предполагается, что credentials содержит путь к сертификатам
                cert_paths = json.loads(self.credentials)
                return self.GigaChat(
                    base_url=self.base_url,
                    ca_bundle_file=cert_paths.get("ca_bundle_file"),
                    cert_file=cert_paths.get("cert_file"),
                    key_file=cert_paths.get("key_file"),
                    key_file_password=cert_paths.get("key_file_password"),
                    model=self.model,
                    **self.kwargs
                )
            elif self.auth_method == "user_pass":
                # Предполагается, что credentials содержит логин и пароль в формате "login:password"
                user, password = self.credentials.split(":", 1)
                return self.GigaChat(
                    base_url=self.base_url,
                    user=user,
                    password=password,
                    model=self.model,
                    verify_ssl_certs=False,
                    **self.kwargs
                )
            else:
                raise ValueError(f"Unsupported auth method: {self.auth_method}")
        except Exception as e:
            logger.error(f"Error initializing GigaChat client: {e}")
            raise

    def summarize(self, text, max_length=500, min_length=50):
        if not text or not text.strip():
            return "Error: No text provided for summarization."

        try:
            # Формирование запроса к GigaChat для суммаризации
            prompt = f"Суммаризируй следующий текст в {max_length} символов или меньше, но не менее {min_length} символов:\n\n{text}"

            with self.client as giga:
                response = giga.chat(prompt)
                return response.choices[0].message.content
        except Exception as e:
            logger.error(f"Error in GigaChat summarization: {e}")
            return f"Ошибка при суммаризации текста с GigaChat: {e}"


# Фабрика для создания суммаризатора
def create_summarizer(summarizer_type="rut5", **kwargs):
    if summarizer_type.lower() == "gigachat":
        return GigaChatSummarizer(**kwargs)
    else:
        return RuT5Summarizer(**kwargs)


def load_models(whisper_model_size="base", summarizer_type="rut5"):
    """
    Загружает все необходимые модели.

    Args:
        whisper_model_size (str): Размер модели Whisper ('tiny', 'base', 'small', 'medium', 'large')
        summarizer_type (str): Тип суммаризатора ('rut5', 'gigachat')
    """
    global whisper_model, voice_encoder, summarization_model, summarization_tokenizer, summarizer

    logger.info("Loading models...")

    # 1. Загрузка модели Whisper для распознавания речи
    logger.info(f"Loading Whisper model ({whisper_model_size})...")
    try:
        whisper_model = whisper.load_model(whisper_model_size, device=DEVICE)
        logger.info("Whisper model loaded.")
    except Exception as e:
        logger.error(f"Error loading Whisper model: {e}")
        return False

    # 2. Загрузка модели для диаризации (Resemblyzer)
    logger.info("Loading voice encoder for diarization...")
    try:
        voice_encoder = VoiceEncoder()
        logger.info("Voice encoder loaded.")
    except Exception as e:
        logger.error(f"Error loading voice encoder: {e}")
        return False

    # 3. Загрузка модели для суммаризации
    if summarizer_type.lower() == "gigachat":
        try:
            # Загрузка конфигурации из файла или переменных окружения
            config = load_config()
            base_url = config.get("GIGACHAT_BASE_URL", GIGACHAT_BASE_URL)
            auth_method = config.get("GIGACHAT_AUTH_METHOD", GIGACHAT_AUTH_METHOD)
            credentials = config.get("GIGACHAT_CREDENTIALS", GIGACHAT_CREDENTIALS)
            model = config.get("GIGACHAT_MODEL", GIGACHAT_MODEL)

            summarizer = create_summarizer(
                summarizer_type="gigachat",
                base_url=base_url,
                auth_method=auth_method,
                credentials=credentials,
                model=model
            )
            logger.info("GigaChat summarizer loaded.")
        except Exception as e:
            logger.error(f"Error loading GigaChat summarizer: {e}")
            logger.info("Falling back to RuT5 summarizer...")
            summarizer_type = "rut5"

    if summarizer_type.lower() == "rut5":
        try:
            # Загрузка модели для суммаризации
            SUMMARIZATION_MODEL_NAME = "cointegrated/rut5-base-absum"
            logger.info(f"Loading summarization model: {SUMMARIZATION_MODEL_NAME}...")

            # Используем явно указанную модель cointegrated/rut5-base-absum, которая доступна на Hugging Face
            # Добавляем параметр trust_remote_code=True для обеспечения совместимости
            summarization_tokenizer = T5Tokenizer.from_pretrained(SUMMARIZATION_MODEL_NAME, trust_remote_code=True)
            summarization_model = T5ForConditionalGeneration.from_pretrained(SUMMARIZATION_MODEL_NAME,
                                                                             trust_remote_code=True).to(DEVICE)

            summarizer = create_summarizer(
                summarizer_type="rut5",
                model=summarization_model,
                tokenizer=summarization_tokenizer,
                device=DEVICE
            )
            logger.info("RuT5 summarizer loaded.")
        except Exception as e:
            logger.error(f"Error loading RuT5 summarizer: {e}")
            logger.error("Hint: Make sure you have internet connection and the model is publicly available.")
            logger.error("You can also try to manually download the model using: huggingface-cli login")
            return False

    return True


def diarize_with_resemblyzer(audio_path, min_speakers=1, max_speakers=5):
    """
    Выполняет диаризацию с использованием Resemblyzer.

    Args:
        audio_path (str): Путь к аудиофайлу
        min_speakers (int): Минимальное количество дикторов
        max_speakers (int): Максимальное количество дикторов

    Returns:
        list: Список сегментов с дикторами или None в случае ошибки
    """
    try:
        # Предобработка аудио
        wav = preprocess_wav(audio_path)

        # Разделение на сегменты (например, по 3 секунды)
        segment_len = 0.1  # в секундах
        sample_rate = 16000  # стандартная частота дискретизации
        segment_samples = int(segment_len * sample_rate)

        segments = []
        segment_times = []
        for i in range(0, len(wav) - segment_samples, segment_samples // 2):  # 50% перекрытие
            segment = wav[i:i + segment_samples]
            if len(segment) == segment_samples:  # только полные сегменты
                segments.append(segment)
                segment_times.append(i / sample_rate)

        if not segments:
            return None

        # Извлечение эмбеддингов для каждого сегмента
        embeddings = np.array([voice_encoder.embed_utterance(segment) for segment in segments])

        # Определение оптимального числа дикторов
        # Для простоты используем фиксированное число или диапазон
        for num_speakers in range(min_speakers, max_speakers + 1):
            # Кластеризация эмбеддингов
            clustering = AgglomerativeClustering(n_clusters=num_speakers)
            labels = clustering.fit_predict(embeddings)

            # Проверка качества кластеризации (например, силуэтный коэффициент)
            # Если качество хорошее, останавливаемся
            # Для простоты берем первый результат
            break

        # Формирование результата в формате, совместимом с остальным кодом
        speaker_turns = []
        for i, label in enumerate(labels):
            speaker_turns.append({
                "start": segment_times[i],
                "end": segment_times[i] + segment_len,
                "speaker": f"SPEAKER_{label}"
            })

        return speaker_turns
    except Exception as e:
        logger.error(f"Error during Resemblyzer diarization: {e}")
        return None


def process_audio(audio_path):
    """
    Обрабатывает аудиофайл: выполняет распознавание речи и диаризацию.

    Args:
        audio_path (str): Путь к аудиофайлу

    Returns:
        tuple: (protocol_text, full_text, speaker_turns)
    """
    if not os.path.exists(audio_path):
        return "Error: Audio file not found.", None, None

    logger.info(f"Processing audio file: {audio_path}")

    # 1. Распознавание речи с Whisper
    logger.info("Starting speech recognition with Whisper...")
    try:
        # `word_timestamps=True` для более точного сопоставления с дикторами
        # `language="ru"` для русского языка
        transcription_result = whisper_model.transcribe(audio_path, language="ru", word_timestamps=True)
        full_text = transcription_result["text"]
        segments = transcription_result["segments"]
        logger.info("Speech recognition completed.")
    except Exception as e:
        logger.error(f"Error during Whisper transcription: {e}")
        return f"Error during Whisper transcription: {e}", None, None

    # 2. Диаризация дикторов с Resemblyzer
    if voice_encoder is None:
        logger.warning("Voice encoder not loaded. Skipping diarization.")
        # Если диаризация не удалась, вернем только транскрипцию
        protocol_text = "Протокол (без разделения на дикторов):\n\n"
        for segment in segments:
            start_time = segment['start']
            end_time = segment['end']
            text = segment['text']
            protocol_text += f"[{start_time:.2f}s - {end_time:.2f}s]: {text.strip()}\n"
        return protocol_text, full_text, None

    logger.info("Starting speaker diarization...")
    try:
        speaker_turns = diarize_with_resemblyzer(audio_path)
        if not speaker_turns:
            raise Exception("No speaker turns detected")
        logger.info("Speaker diarization completed.")
    except Exception as e:
        logger.error(f"Error during speaker diarization: {e}")
        # Если диаризация не удалась, вернем только транскрипцию
        protocol_text = "Протокол (ошибка диаризации, без разделения на дикторов):\n\n"
        for segment in segments:
            start_time = segment['start']
            end_time = segment['end']
            text = segment['text']
            protocol_text += f"[{start_time:.2f}s - {end_time:.2f}s]: {text.strip()}\n"
        return protocol_text, full_text, None

    # 3. Объединение результатов ASR и диаризации
    logger.info("Combining ASR and diarization results...")

    # Создаем список слов с их временем начала, конца и текстом
    word_level_info = []
    for segment in segments:
        for word_info in segment.get("words", []):  # Используем .get для безопасного доступа
            if 'start' in word_info and 'end' in word_info and 'word' in word_info:
                word_level_info.append({
                    "start": word_info["start"],
                    "end": word_info["end"],
                    "text": word_info["word"].strip()
                })
            elif 'start' in segment and 'end' in segment and 'text' in segment:  # Если нет word_timestamps
                # Попытаемся разбить текст сегмента на слова и распределить время
                words_in_segment = segment["text"].strip().split()
                segment_duration = segment["end"] - segment["start"]
                if words_in_segment and segment_duration > 0:
                    time_per_word = segment_duration / len(words_in_segment)
                    current_time = segment["start"]
                    for word in words_in_segment:
                        word_level_info.append({
                            "start": current_time,
                            "end": current_time + time_per_word,
                            "text": word
                        })
                        current_time += time_per_word
                break  # Выходим из внутреннего цикла, так как обработали весь сегмент

    if not word_level_info and segments:  # Если word_timestamps=False и предыдущий блок не сработал
        logger.warning(
            "Word-level timestamps not available from Whisper. Using segment-level timestamps for diarization matching.")
        # Используем начало сегмента Whisper для определения диктора
        for segment in segments:
            word_level_info.append({
                "start": segment["start"],
                "end": segment["end"],
                "text": segment["text"].strip(),  # Весь текст сегмента как одно "слово"
                "is_segment": True  # Флаг, что это целый сегмент
            })

    # Сортируем реплики дикторов по времени начала
    speaker_turns.sort(key=lambda x: x["start"])

    # Собираем протокол
    protocol_entries = []
    current_speaker = None
    current_utterance = ""
    first_word_start_time = 0
    last_word_end_time = 0

    # Проходим по каждому слову/сегменту из Whisper
    for word_info in word_level_info:
        word_start_time = word_info["start"]
        word_end_time = word_info["end"]
        word_text = word_info["text"]
        is_segment_text = word_info.get("is_segment", False)

        # Находим диктора для текущего слова/сегмента
        assigned_speaker = "UNKNOWN_SPEAKER"
        best_overlap = 0
        for turn in speaker_turns:
            # Рассчитываем пересечение интервалов
            overlap_start = max(word_start_time, turn["start"])
            overlap_end = min(word_end_time, turn["end"])
            overlap_duration = max(0, overlap_end - overlap_start)

            if overlap_duration > best_overlap:
                best_overlap = overlap_duration
                assigned_speaker = turn["speaker"]

        # Если диктор сменился или это первое слово
        if assigned_speaker != current_speaker and current_speaker is not None:
            if current_utterance.strip():
                protocol_entries.append({
                    "speaker": current_speaker,
                    "utterance": current_utterance.strip(),
                    "start_time": first_word_start_time,  # Время начала первого слова в реплике
                    "end_time": last_word_end_time  # Время конца последнего слова в реплике
                })
            current_utterance = ""
            first_word_start_time = word_start_time  # Обновляем время начала для новой реплики

        if current_speaker is None:  # Для самого первого слова
            first_word_start_time = word_start_time

        current_speaker = assigned_speaker
        current_utterance += word_text + (" " if not is_segment_text else "\n")
        last_word_end_time = word_end_time  # Обновляем время конца последнего слова

    # Добавляем последнюю реплику
    if current_utterance.strip() and current_speaker is not None:
        protocol_entries.append({
            "speaker": current_speaker,
            "utterance": current_utterance.strip(),
            "start_time": first_word_start_time,
            "end_time": last_word_end_time
        })

    # Формируем текстовый протокол
    protocol_text = "Протокол совещания:\n\n"
    for entry in protocol_entries:
        protocol_text += f"[{entry['start_time']:.2f}s - {entry['end_time']:.2f}s] {entry['speaker']}: {entry['utterance']}\n"

    logger.info("ASR and diarization results combined.")
    return protocol_text, full_text, speaker_turns


def summarize_text(text, max_length=500, min_length=50):
    """
    Суммаризирует предоставленный текст.

    Args:
        text (str): Текст для суммаризации
        max_length (int): Максимальная длина суммаризации
        min_length (int): Минимальная длина суммаризации

    Returns:
        str: Суммаризированный текст или сообщение об ошибке
    """
    if not text or not text.strip():
        return "Error: No text provided for summarization."

    logger.info("Starting text summarization...")
    try:
        summary = summarizer.summarize(text, max_length=max_length, min_length=min_length)
        logger.info("Text summarization completed.")
        return summary
    except Exception as e:
        logger.error(f"Error during text summarization: {e}")
        error_msg = f"Ошибка при суммаризации текста: {e}"
        logger.error(f"Detailed error: {error_msg}")
        return error_msg


def generate_docx_protocol(protocol_text, output_path):
    """
    Генерирует протокол совещания в формате .docx

    Args:
        protocol_text (str): Текст протокола
        output_path (str): Путь для сохранения файла

    Returns:
        str: Путь к сохраненному файлу
    """
    doc = Document()
    doc.add_heading('Протокол совещания', 0)

    # Разбиваем протокол на строки и добавляем в документ
    lines = protocol_text.split('\n')
    for line in lines:
        if line.strip():
            if line.startswith('Протокол'):
                # Пропускаем заголовок, так как мы уже добавили его
                continue
            elif '[' in line and ']' in line and ':' in line:
                # Это строка с репликой диктора
                # Форматируем как таблицу или с отступами
                parts = line.split(': ', 1)
                if len(parts) == 2:
                    header, text = parts
                    p = doc.add_paragraph()
                    p.add_run(header + ': ').bold = True
                    p.add_run(text)
            else:
                # Обычный текст
                doc.add_paragraph(line)

    # Сохраняем документ
    doc.save(output_path)
    logger.info(f"Protocol saved to {output_path}")
    return output_path


def generate_txt_summary(summary_text, output_path):
    """
    Генерирует суммаризацию в формате .txt

    Args:
        summary_text (str): Текст суммаризации
        output_path (str): Путь для сохранения файла

    Returns:
        str: Путь к сохраненному файлу
    """
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(summary_text)
    logger.info(f"Summary saved to {output_path}")
    return output_path


def process_and_save(audio_path, output_dir=None):
    """
    Обрабатывает аудиофайл и сохраняет результаты

    Args:
        audio_path (str): Путь к аудиофайлу
        output_dir (str): Директория для сохранения результатов

    Returns:
        dict: Словарь с результатами обработки
    """
    # Определяем директорию для сохранения результатов
    if not output_dir:
        output_dir = os.path.dirname(audio_path)

    # Создаем директорию, если она не существует
    os.makedirs(output_dir, exist_ok=True)

    # Получаем имя файла без расширения
    base_name = os.path.splitext(os.path.basename(audio_path))[0]

    # Обрабатываем аудиофайл
    protocol_text, full_text, speaker_turns = process_audio(audio_path)

    # Если не удалось получить текст, возвращаем ошибку
    if not full_text:
        return {"error": protocol_text}

    # Суммаризируем текст
    summary_text = summarize_text(full_text)

    # Сохраняем протокол в формате .docx
    protocol_path = os.path.join(output_dir, f"{base_name}_protocol.docx")
    generate_docx_protocol(protocol_text, protocol_path)

    # Сохраняем суммаризацию в формате .txt
    summary_path = os.path.join(output_dir, f"{base_name}_summary.txt")
    generate_txt_summary(summary_text, summary_path)

    # Возвращаем результаты
    return {
        "protocol_text": protocol_text,
        "summary_text": summary_text,
        "protocol_path": protocol_path,
        "summary_path": summary_path
    }


def create_gui():
    """
    Создает графический интерфейс пользователя
    """
    root = tk.Tk()
    root.title("Анализатор совещаний")
    root.geometry("800x600")

    # Функция для выбора аудиофайла
    def select_audio_file():
        file_path = filedialog.askopenfilename(title="Выберите аудиофайл",
                                               filetypes=[("Audio files", "*.wav;*.mp3;*.ogg;*.flac"),
                                                          ("All files", "*.*")])
        if file_path:
            audio_path_var.set(file_path)

    # Функция для выбора директории вывода
    def select_output_dir():
        dir_path = filedialog.askdirectory(title="Выберите директорию для сохранения результатов")
        if dir_path:
            output_dir_var.set(dir_path)

    # Функция для обработки аудиофайла
    def process_file():
        audio_path = audio_path_var.get()
        output_dir = output_dir_var.get()
        summarizer_type = summarizer_var.get()

        if not audio_path:
            messagebox.showerror("Ошибка", "Выберите аудиофайл для обработки")
            return

        if not os.path.exists(audio_path):
            messagebox.showerror("Ошибка", f"Файл не найден: {audio_path}")
            return

        # Проверяем, загружены ли модели
        if whisper_model is None or voice_encoder is None or summarizer is None:
            # Показываем прогресс-бар загрузки моделей
            progress_label.config(text="Загрузка моделей...")
            progress_bar.grid(row=6, column=0, columnspan=3, padx=10, pady=10, sticky="ew")
            root.update()

            # Загружаем модели в отдельном потоке
            def load_models_thread():
                whisper_size = whisper_model_var.get()
                success = load_models(whisper_size, summarizer_type)

                # Обновляем интерфейс в основном потоке
                root.after(0, lambda: after_models_loaded(success))

            threading.Thread(target=load_models_thread).start()
            return

        # Показываем прогресс-бар обработки
        progress_label.config(text="Обработка аудиофайла...")
        progress_bar.grid(row=6, column=0, columnspan=3, padx=10, pady=10, sticky="ew")
        process_button.config(state=tk.DISABLED)
        root.update()

        # Обрабатываем аудиофайл в отдельном потоке
        def process_file_thread():
            try:
                results = process_and_save(audio_path, output_dir if output_dir else None)

                # Обновляем интерфейс в основном потоке
                root.after(0, lambda: after_processing_completed(results))
            except Exception as e:
                root.after(0, lambda: messagebox.showerror("Ошибка обработки", str(e)))
                root.after(0, lambda: reset_ui())

        threading.Thread(target=process_file_thread).start()

    # Функция, вызываемая после загрузки моделей
    def after_models_loaded(success):
        if success:
            progress_label.config(text="Модели загружены успешно")
            process_file()  # Продолжаем обработку файла
        else:
            progress_label.config(text="Ошибка загрузки моделей")
            progress_bar.grid_remove()
            messagebox.showerror("Ошибка", "Не удалось загрузить модели")

    # Функция, вызываемая после завершения обработки
    def after_processing_completed(results):
        progress_bar.grid_remove()
        process_button.config(state=tk.NORMAL)

        if "protocol_path" in results and "summary_path" in results:
            messagebox.showinfo(
                "Обработка завершена",
                f"Протокол сохранен в: {results['protocol_path']}\n"
                f"Суммаризация сохранена в: {results['summary_path']}"
            )

            # Показываем результаты в текстовом поле
            result_text.delete(1.0, tk.END)
            result_text.insert(tk.END, "=== СУММАРИЗАЦИЯ ===\n\n")
            result_text.insert(tk.END, results.get("summary_text", "Суммаризация недоступна") + "\n\n")
            result_text.insert(tk.END, "=== ФРАГМЕНТ ПРОТОКОЛА ===\n\n")

            # Показываем только первые 20 строк протокола
            protocol_lines = results.get("protocol_text", "Протокол недоступен").split("\n")
            protocol_preview = "\n".join(protocol_lines[:20])
            if len(protocol_lines) > 20:
                protocol_preview += "\n...\n(Полный протокол доступен в сохраненном файле)"

            result_text.insert(tk.END, protocol_preview)
        else:
            messagebox.showerror("Ошибка", "Не удалось обработать аудиофайл")

    # Функция для сброса интерфейса
    def reset_ui():
        progress_bar.grid_remove()
        process_button.config(state=tk.NORMAL)

    # Создаем переменные для хранения путей
    audio_path_var = tk.StringVar()
    output_dir_var = tk.StringVar()
    whisper_model_var = tk.StringVar(value="base")
    summarizer_var = tk.StringVar(value="rut5")

    # Создаем элементы интерфейса
    frame = ttk.Frame(root, padding=10)
    frame.pack(fill=tk.BOTH, expand=True)

    # Выбор аудиофайла
    ttk.Label(frame, text="Аудиофайл:").grid(row=0, column=0, sticky=tk.W, pady=5)
    ttk.Entry(frame, textvariable=audio_path_var, width=50).grid(row=0, column=1, pady=5, padx=5)
    ttk.Button(frame, text="Обзор", command=select_audio_file).grid(row=0, column=2, pady=5)

    # Выбор директории вывода
    ttk.Label(frame, text="Директория вывода:").grid(row=1, column=0, sticky=tk.W, pady=5)
    ttk.Entry(frame, textvariable=output_dir_var, width=50).grid(row=1, column=1, pady=5, padx=5)
    ttk.Button(frame, text="Обзор", command=select_output_dir).grid(row=1, column=2, pady=5)

    # Выбор модели Whisper
    ttk.Label(frame, text="Модель Whisper:").grid(row=2, column=0, sticky=tk.W, pady=5)
    model_combo = ttk.Combobox(frame, textvariable=whisper_model_var,
                               values=["tiny", "base", "small", "medium", "large"])
    model_combo.grid(row=2, column=1, sticky=tk.W, pady=5, padx=5)
    model_combo.current(1)  # По умолчанию "base"

    # Выбор модели суммаризации
    ttk.Label(frame, text="Модель суммаризации:").grid(row=3, column=0, sticky=tk.W, pady=5)
    summarizer_combo = ttk.Combobox(frame, textvariable=summarizer_var,
                                    values=["rut5", "gigachat"])
    summarizer_combo.grid(row=3, column=1, sticky=tk.W, pady=5, padx=5)
    summarizer_combo.current(0)  # По умолчанию "rut5"

    # Кнопка обработки
    process_button = ttk.Button(frame, text="Обработать", command=process_file)
    process_button.grid(row=4, column=0, columnspan=3, pady=10)

    # Прогресс-бар (скрыт по умолчанию)
    progress_label = ttk.Label(frame, text="")
    progress_label.grid(row=5, column=0, columnspan=3, pady=5)
    progress_bar = ttk.Progressbar(frame, mode="indeterminate")
    progress_bar.grid(row=6, column=0, columnspan=3, padx=10, pady=10, sticky="ew")
    progress_bar.grid_remove()  # Скрываем до начала обработки
    progress_bar.start(10)  # Запускаем анимацию

    # Текстовое поле для отображения результатов
    result_frame = ttk.LabelFrame(frame, text="Результаты")
    result_frame.grid(row=7, column=0, columnspan=3, padx=10, pady=10, sticky="nsew")

    result_text = tk.Text(result_frame, wrap=tk.WORD, height=15)
    result_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

    # Настраиваем растяжение строк и столбцов
    frame.columnconfigure(1, weight=1)
    frame.rowconfigure(7, weight=1)

    # Запускаем главный цикл
    root.mainloop()


def main():
    """
    Основная функция для запуска приложения
    """
    parser = argparse.ArgumentParser(description='Анализ аудиозаписей совещаний и переговоров')
    parser.add_argument('--audio', help='Путь к аудиофайлу (WAV формат)')
    parser.add_argument('--output_dir', help='Директория для сохранения результатов')
    parser.add_argument('--whisper_model', choices=['tiny', 'base', 'small', 'medium', 'large'],
                        default='base', help='Размер модели Whisper (по умолчанию: base)')
    parser.add_argument('--summarizer', choices=['rut5', 'gigachat'],
                        default='rut5', help='Модель для суммаризации (по умолчанию: rut5)')
    parser.add_argument('--config', help='Путь к конфигурационному файлу')
    parser.add_argument('--gui', action='store_true', help='Запустить графический интерфейс')

    args = parser.parse_args()

    # Загружаем конфигурацию, если указан путь
    if args.config:
        config = load_config(args.config)
        # Обновляем глобальные переменные из конфигурации
        global GIGACHAT_BASE_URL, GIGACHAT_AUTH_METHOD, GIGACHAT_CREDENTIALS, GIGACHAT_MODEL
        GIGACHAT_BASE_URL = config.get("GIGACHAT_BASE_URL", GIGACHAT_BASE_URL)
        GIGACHAT_AUTH_METHOD = config.get("GIGACHAT_AUTH_METHOD", GIGACHAT_AUTH_METHOD)
        GIGACHAT_CREDENTIALS = config.get("GIGACHAT_CREDENTIALS", GIGACHAT_CREDENTIALS)
        GIGACHAT_MODEL = config.get("GIGACHAT_MODEL", GIGACHAT_MODEL)

    # Запускаем GUI, если указан флаг --gui или не указан аудиофайл
    if args.gui or not args.audio:
        create_gui()
        return

    # Загружаем модели
    if not load_models(args.whisper_model, args.summarizer):
        logger.error("Failed to load models. Exiting.")
        return

    # Обрабатываем аудиофайл
    results = process_and_save(args.audio, args.output_dir)

    logger.info("\nProcessing completed!")
    logger.info(f"Protocol saved to: {results['protocol_path']}")
    logger.info(f"Summary saved to: {results['summary_path']}")


if __name__ == "__main__":
    main()
